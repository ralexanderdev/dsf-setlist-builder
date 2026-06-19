import streamlit as st
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
import io
import json

# Your complete master list data categorized by list type
MASTER_SONGS = [
    ("Intro Jam (General)", "04:00", "A"), ("Good Times Roll", "03:05", "A"),
    ("Affirmation lover", "03:44", "A"), ("Barriers", "04:24", "A"),
    ("Hydroplane", "04:50", "A"), ("Side 2 Side", "03:37", "A"),
    ("Go Away", "05:38", "A"), ("Don't Be Anybody Else", "04:08", "A"),
    ("Show me the way", "03:56", "A"), ("What you wanna do", "03:02", "A"),
    ("Can you feel it", "04:00", "B"), ("Lets Go", "03:27", "B"),
    ("Come Down", "04:00", "B"), ("Night is Young", "03:12", "B"),
    ("Neo Soul", "04:40", "C"), ("Move Baby", "04:16", "C"),
    ("Beleive", "03:44", "C"), ("Life is Funk", "03:28", "C"),
    ("Treat Love Simple", "04:16", "C"), ("Levitate", "04:02", "C"),
    ("Red Light", "04:21", "C"), ("Seven Nation Army", "04:01", "C"),
    ("Pull Yourself Together", "03:55", "Intro"), ("Intro (Lets Go Jam)", "00:55", "Intro")
]

# Set up browser page settings
st.set_page_config(page_title="Band Setlist Builder", layout="wide")
st.title("🎸 Band Setlist Builder")

# ---------------- TOP PANEL: GIG METADATA ----------------
col1, col2 = st.columns(2)
with col1:
    gig_name = st.text_input("Gig / Venue Name", "Silverburn Festival")
with col2:
    gig_date = st.date_input("Gig Date", datetime.today())

# Filter master data into tab buckets
a_tracks = [s for s in MASTER_SONGS if s[2] == "A"]
b_tracks = [s for s in MASTER_SONGS if s[2] == "B"]
c_intros = [s for s in MASTER_SONGS if s[2] in ["C", "Intro"]]

# Setup side-by-side workspace columns
left_col, right_col = st.columns(2)

# ---------------- LEFT COLUMN: SPLIT SONG POOLS ----------------
with left_col:
    st.subheader("🎵 Song Pool Categories")
    st.caption("Check the boxes of the songs you want to make available for tonight's gig:")
    
    tab1, tab2, tab3 = st.tabs(["A List", "B List", "C List / Intros"])
    
    with tab1:
        selected_a = []
        for s in a_tracks:
            label_str = f"{s[0]} ({s[1]})"
            if st.checkbox(label_str, key=f"check_{s[0]}"):
                selected_a.append(label_str)
                
    with tab2:
        selected_b = []
        for s in b_tracks:
            label_str = f"{s[0]} ({s[1]})"
            if st.checkbox(label_str, key=f"check_{s[0]}"):
                selected_b.append(label_str)
                
    with tab3:
        selected_c = []
        for s in c_intros:
            label_str = f"{s[0]} ({s[1]})"
            if st.checkbox(label_str, key=f"check_{s[0]}"):
                selected_c.append(label_str)

# Combine active checked tracks
chosen_items = selected_a + selected_b + selected_c

# Track state of list order across updates
if "current_order" not in st.session_state or set(st.session_state.current_order) != set(chosen_items):
    st.session_state.current_order = chosen_items.copy()

# ---------------- RIGHT COLUMN: PIXEL-PERFECT LISTBOX ----------------
with right_col:
    st.markdown("**Active Setlist Order** (Drag items vertically to reorder)")
    
    if not chosen_items:
        st.info("Check songs from the left panels to populate your setlist.")
        final_setlist_names = []
    else:
        # Read incoming sorted data from the HTML component
        query_order = st.query_params.get("order", None)
        if query_order:
            try:
                parsed_order = json.loads(query_order)
                # Ensure the parsed items still accurately match current pool selections
                if set(parsed_order) == set(chosen_items):
                    st.session_state.current_order = parsed_order
            except:
                pass

        # HTML/JS Code injection to precisely emulate the image layout
        escaped_items = json.dumps(st.session_state.current_order)
        
        listbox_html = f"""
        <script src="https://cdnjs.cloudflare.com/ajax/libs/Sortable/1.15.0/Sortable.min.js"></script>
        <style>
            body {{ margin: 0; font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif; }}
            .listbox {{
                border: 1px solid #999;
                background: white;
                height: 260px;
                overflow-y: auto;
                user-select: none;
            }}
            .list-item {{
                padding: 4px 8px;
                font-size: 14px;
                color: black;
                cursor: default;
                white-space: nowrap;
            }}
            .list-item:hover {{
                background-color: #f0f0f0;
            }}
            /* Matches the exact deep blue selection band seen in the screenshot */
            .list-item.sortable-chosen, .list-item:active, .list-item.selected-active {{
                background-color: #0056b3 !important;
                color: white !important;
            }}
        </style>
        
        <div id="items-container" class="listbox"></div>

        <script>
            const items = {escaped_items};
            const container = document.getElementById('items-container');
            
            items.forEach((item, index) => {{
                const div = document.createElement('div');
                div.className = 'list-item';
                div.textContent = item;
                // Add highlight on click mimicking standard selection lists
                div.addEventListener('mousedown', function() {{
                    document.querySelectorAll('.list-item').forEach(el => el.classList.remove('selected-active'));
                    div.classList.add('selected-active');
                }});
                container.appendChild(div);
            }});

            Sortable.create(container, {{
                animation: 150,
                onEnd: function() {{
                    const updatedOrder = [];
                    document.querySelectorAll('.list-item').forEach(el => {{
                        updatedOrder.push(el.textContent);
                    }});
                    
                    // Send back to Streamlit URL query parameter stack safely
                    const url = new URL(window.parent.location.href);
                    url.searchParams.set('order', JSON.stringify(updatedOrder));
                    window.parent.location.href = url.href;
                }}
            }});
        </script>
        """
        
        # Render the custom view component frame
        st.components.v1.html(listbox_html, height=270)
        final_setlist_names = st.session_state.current_order

        # Live Performance Running Time Calculations
        total_secs = 0
        for item in final_setlist_names:
            try:
                time_part = item.split('(')[-1].replace(')', '').strip()
                mins, secs = map(int, time_part.split(':'))
                total_secs += mins * 60 + secs
            except:
                pass
        
        hours, remainder = divmod(total_secs, 3600)
        minutes, seconds = divmod(remainder, 60)
        time_str = f"{hours:02d}:{minutes:02d}:{seconds:02d}" if hours > 0 else f"{minutes:02d}:{seconds:02d}"
        
        # Performance Summary Panels
        st.markdown("---")
        m_col1, m_col2 = st.columns(2)
        with m_col1:
            st.metric(label="Total Tracks", value=len(final_setlist_names))
        with m_col2:
            st.metric(label="Total Runtime", value=time_str)

# ---------------- WORD DOC GENERATION ----------------
        doc = Document()
        
        # 1-inch printable stage margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Main Heading: Bold and Underlined (Matches WhatsApp Image 2026-05-31 at 12.31.45.jpg)
        title = doc.add_paragraph()
        title_run = title.add_run(f"Dopesickfly - {gig_name} {gig_date.strftime('%d%m%y')}")
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.underline = True
        title.paragraph_format.space_after = Pt(24) # Gives a nice gap after the header

        # Print clean, sequentially numbered list items without round bullets
        for idx, item in enumerate(final_setlist_names, start=1):
            # Strip out the trailing runtime (04:00) so it's just the clean song name
            display_name = item.split('(')[0].strip()
            
            p = doc.add_paragraph()
            # Tight line spacing to match the reference document look
            p.paragraph_format.space_after = Pt(6) 
            p.paragraph_format.line_spacing = 1.15
            
            # Format as: "1.   Song Name"
            run = p.add_run(f"{idx}.   {display_name}")
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            run.font.bold = True

        # Render file compiler payload in-memory
        bio = io.BytesIO()
        doc.save(bio)
        
        # Streamlit safe download deployment trigger
        st.download_button(
            label="💾 Download Stage Word Doc (.docx)",
            data=bio.getvalue(),
            file_name=f"Dopesickfly - {gig_name} {gig_date.strftime('%Y-%m-%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )